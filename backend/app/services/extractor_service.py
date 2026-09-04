import io
import json
import csv
import re
from typing import Dict, Any, List, Tuple, Optional
from pypdf import PdfReader
import docx

NO_EXTRACT_MSG = "Could not extract this field from uploaded evidence."

class DocumentExtractorService:
    """
    Service for extracting raw text, structured fields, and page provenance mapping
    from uploaded export documents (PDF, DOCX, TXT, CSV, JSON).
    """

    def extract_document(self, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        
        pages_text: List[Tuple[str, str]] = []  # List of (page_label, page_text)
        full_text = ""

        if ext == "pdf":
            pages_text, full_text = self._extract_pdf(content_bytes)
        elif ext == "docx":
            pages_text, full_text = self._extract_docx(content_bytes)
        elif ext in ["txt", "log"]:
            pages_text, full_text = self._extract_txt(content_bytes)
        elif ext == "csv":
            pages_text, full_text = self._extract_csv(content_bytes)
        elif ext == "json":
            pages_text, full_text = self._extract_json(content_bytes)
        else:
            pages_text, full_text = self._extract_txt(content_bytes)

        # Extract structured fields with provenance tracking
        extracted_fields, provenance_map = self._extract_structured_fields(filename, full_text, pages_text)

        return {
            "full_text": full_text,
            "pages_count": len(pages_text),
            "pages": [{"page_label": p[0], "text": p[1]} for p in pages_text],
            "extracted_fields": extracted_fields,
            "provenance_map": provenance_map,
            "file_format": ext.upper()
        }

    def _extract_pdf(self, content_bytes: bytes) -> Tuple[List[Tuple[str, str]], str]:
        pages_text = []
        full_text = ""
        try:
            reader = PdfReader(io.BytesIO(content_bytes))
            for idx, page in enumerate(reader.pages):
                page_label = f"Page {idx + 1}"
                txt = page.extract_text() or ""
                pages_text.append((page_label, txt))
                full_text += f"\n--- {page_label} ---\n" + txt
        except Exception as e:
            full_text = f"PDF Parsing Error: {str(e)}"
            pages_text = [("Page 1", full_text)]
        return pages_text, full_text

    def _extract_docx(self, content_bytes: bytes) -> Tuple[List[Tuple[str, str]], str]:
        pages_text = []
        full_text = ""
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            full_text = "\n".join(paragraphs)
            pages_text = [("Section 1", full_text)]
        except Exception as e:
            full_text = f"DOCX Parsing Error: {str(e)}"
            pages_text = [("Section 1", full_text)]
        return pages_text, full_text

    def _extract_txt(self, content_bytes: bytes) -> Tuple[List[Tuple[str, str]], str]:
        try:
            full_text = content_bytes.decode("utf-8")
        except Exception:
            full_text = content_bytes.decode("latin-1", errors="ignore")
        return [("Page 1", full_text)], full_text

    def _extract_csv(self, content_bytes: bytes) -> Tuple[List[Tuple[str, str]], str]:
        try:
            text = content_bytes.decode("utf-8")
        except Exception:
            text = content_bytes.decode("latin-1", errors="ignore")
        
        lines = []
        reader = csv.reader(io.StringIO(text))
        for row_idx, row in enumerate(reader):
            lines.append(f"Row {row_idx + 1}: " + " | ".join(row))
        full_text = "\n".join(lines)
        return [("CSV File", full_text)], full_text

    def _extract_json(self, content_bytes: bytes) -> Tuple[List[Tuple[str, str]], str]:
        try:
            raw_str = content_bytes.decode("utf-8")
            data = json.loads(raw_str)
            full_text = json.dumps(data, indent=2)
        except Exception:
            data = None
            full_text = content_bytes.decode("utf-8", errors="ignore")
        return [("JSON Document", full_text)], full_text

    def _find_page_for_text(self, target_str: str, pages_text: List[Tuple[str, str]]) -> str:
        if not target_str:
            return "Page 1"
        target_lower = str(target_str).lower()
        for page_label, text in pages_text:
            if target_lower in text.lower():
                return page_label
        return "Page 1"

    def _extract_structured_fields(
        self,
        filename: str,
        full_text: str,
        pages_text: List[Tuple[str, str]]
    ) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        extracted = {}
        provenance = {}

        def set_field(key: str, value: Any, raw_snippet: Optional[str] = None):
            if value is not None and str(value).strip() != "" and str(value).strip() != NO_EXTRACT_MSG:
                extracted[key] = value
                page_label = self._find_page_for_text(raw_snippet or str(value), pages_text)
                provenance[key] = {
                    "value": value,
                    "source_file": filename,
                    "source_page": page_label,
                    "status": "EXTRACTED"
                }
            else:
                extracted[key] = NO_EXTRACT_MSG
                provenance[key] = {
                    "value": NO_EXTRACT_MSG,
                    "source_file": filename,
                    "source_page": "N/A",
                    "status": "MISSING"
                }

        # JSON direct key parsing if JSON document
        try:
            json_obj = json.loads(full_text)
            if isinstance(json_obj, dict):
                if "residue_value" in json_obj and isinstance(json_obj["residue_value"], (int, float)):
                    set_field("residue_value", float(json_obj["residue_value"]), "residue_value")
                if "active_ingredient" in json_obj:
                    set_field("active_ingredient", str(json_obj["active_ingredient"]), "active_ingredient")
                if "product" in json_obj:
                    set_field("product", str(json_obj["product"]), "product")
                if "unit" in json_obj:
                    set_field("unit", str(json_obj["unit"]), "unit")
                if "quantity_kg" in json_obj:
                    set_field("quantity_kg", float(json_obj["quantity_kg"]), "quantity_kg")
        except Exception:
            pass

        # 1. Residue Value extraction
        if "residue_value" not in extracted or extracted["residue_value"] == NO_EXTRACT_MSG:
            res_match = re.search(r'(?:residue|measured|level|concentration)\s*[:=]?\s*([0-9]+\.[0-9]+)\s*(?:mg/kg|ppm)?', full_text, re.IGNORECASE)
            if not res_match:
                res_match = re.search(r'([0-9]+\.[0-9]+)\s*mg/kg', full_text, re.IGNORECASE)
            if res_match:
                set_field("residue_value", float(res_match.group(1)), res_match.group(0))
                set_field("unit", "mg/kg", res_match.group(0))
            else:
                set_field("residue_value", None)
                set_field("unit", None)

        # 2. Active Ingredient
        if "active_ingredient" not in extracted or extracted["active_ingredient"] == NO_EXTRACT_MSG:
            if "imidacloprid" in full_text.lower():
                set_field("active_ingredient", "Imidacloprid", "Imidacloprid")
            elif "buprofezin" in full_text.lower():
                set_field("active_ingredient", "Buprofezin", "Buprofezin")
            elif "chlorpyrifos" in full_text.lower():
                set_field("active_ingredient", "Chlorpyrifos", "Chlorpyrifos")
            else:
                ai_match = re.search(r'(?:active\s*ingredient|active\s*substance|pesticide|chemical)\s*[:=]?\s*([A-Za-z]{3,20})', full_text, re.IGNORECASE)
                if ai_match:
                    set_field("active_ingredient", ai_match.group(1).strip(), ai_match.group(0))
                else:
                    set_field("active_ingredient", None)

        # 3. Quantity KG
        if "quantity_kg" not in extracted or extracted["quantity_kg"] == NO_EXTRACT_MSG:
            qty_match = re.search(r'(?:quantity|weight|net\s*weight)\s*[:=]?\s*([0-9]+(?:\.[0-9]+)?)\s*(?:kg|kgs)?', full_text, re.IGNORECASE)
            if not qty_match:
                qty_match = re.search(r'([0-9]{3,6})\s*kg', full_text, re.IGNORECASE)
            if qty_match:
                set_field("quantity_kg", float(qty_match.group(1)), qty_match.group(0))
            else:
                set_field("quantity_kg", None)

        # 4. Product / Commodity
        if "product" not in extracted or extracted["product"] == NO_EXTRACT_MSG:
            if "mango" in full_text.lower():
                set_field("product", "Mango", "Mango")
            elif "grapes" in full_text.lower():
                set_field("product", "Grapes", "Grapes")
            elif "rice" in full_text.lower():
                set_field("product", "Rice", "Rice")
            elif "papaya" in full_text.lower():
                set_field("product", "Papaya", "Papaya")
            else:
                prod_match = re.search(r'(?:product|commodity|crop)\s*[:=]?\s*([A-Za-z]{3,15})', full_text, re.IGNORECASE)
                if prod_match:
                    set_field("product", prod_match.group(1).strip(), prod_match.group(0))
                else:
                    set_field("product", None)

        # 5. Laboratory Name
        if "laboratory" not in extracted or extracted["laboratory"] == NO_EXTRACT_MSG:
            if "nabl" in full_text.lower():
                set_field("laboratory", "NABL Accredited Export Quality Control Lab", "NABL")
            else:
                lab_match = re.search(r'(?:laboratory|testing\s*lab)\s*[:=]?\s*([A-Za-z0-9\.,\-\s]{4,30})', full_text, re.IGNORECASE)
                if lab_match:
                    set_field("laboratory", lab_match.group(1).strip(), lab_match.group(0))
                else:
                    set_field("laboratory", None)

        # 6. Certificate Number
        if "document_number" not in extracted or extracted["document_number"] == NO_EXTRACT_MSG:
            cert_match = re.search(r'(?:certificate\s*no|invoice\s*no|packing\s*list\s*no|cert\s*no)\s*[:=]?\s*([A-Za-z0-9\-\/]{4,25})', full_text, re.IGNORECASE)
            if cert_match:
                set_field("document_number", cert_match.group(1).strip(), cert_match.group(0))
            else:
                set_field("document_number", None)

        # 7. Batch / Lot Identifier
        if "batch_id" not in extracted or extracted["batch_id"] == NO_EXTRACT_MSG:
            batch_match = re.search(r'(?:batch|lot|sample\s*id)\s*[:=]?\s*([A-Za-z0-9\-\/]{3,20})', full_text, re.IGNORECASE)
            if batch_match:
                set_field("batch_id", batch_match.group(1).strip(), batch_match.group(0))
            else:
                set_field("batch_id", None)

        # 8. Test Result
        if "test_result" not in extracted or extracted["test_result"] == NO_EXTRACT_MSG:
            if re.search(r'\b(PASS|COMPLIANT|ACCEPTED)\b', full_text, re.IGNORECASE):
                set_field("test_result", "PASS", "PASS")
            elif re.search(r'\b(FAIL|NON\-COMPLIANT|REJECTED)\b', full_text, re.IGNORECASE):
                set_field("test_result", "FAIL", "FAIL")
            else:
                set_field("test_result", None)

        return extracted, provenance

extractor_service_instance = DocumentExtractorService()
